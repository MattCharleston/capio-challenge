import pycurl
import certifi
import json
from docx import Document
from docx.text.run import Font, Run
from docx.shared import RGBColor
import datetime
try:
    from StringIO import StringIO
except ImportError:
    from io import StringIO
import unittest

def main(api_key, transaction_id):
	succ, buff = make_request(api_key, transaction_id)
	if (succ != 200):
		req.close()
		return "ERROR"
	parse_response(buff)
	return 0

"""Sends a cURL request using the pycurl library
Returns the response code and the body of the response. 
"""
def make_request(api_key, transaction_id):
	resp = StringIO()
	req = pycurl.Curl()
	url = "https://api.capio.ai/v1/speech/transcript/{0}".format(transaction_id)
	req.setopt(req.URL, url)
	req.setopt(req.CAINFO, certifi.where())
	req.setopt(req.HTTPHEADER, ["apiKey: {0}".format(api_key)])
	req.setopt(req.WRITEFUNCTION, resp.write)
	req.setopt(req.WRITEDATA, resp)
	req.perform()
	resp = resp.getvalue()
	resp_code = req.getinfo(req.RESPONSE_CODE)
	req.close()
	return resp_code, resp

"""Loads the json object and loops through the data adding it to a MS word doc.
"""
def parse_response(body):
	doc = Document()
	document = doc.add_paragraph()
	body = json.loads(body)
	for elm in body:
		result = elm["result"]
		alternative = result[0]
		alternative = alternative['alternative']
		for key in alternative:
			sentance = StringIO()
			time = key['words'][0]['from']
			time = float("{0:.2f}".format(time))
			m, s = divmod(time, 60)
			m = int(m)
			h, m = divmod(m, 60)
			time = "{0}:{1}:{2}".format("%02d" % (h,), "%02d" % (m,), s)
			addto_doc(document, "{}	".format(time), style=True)
			for word in key['words']:
				if word['confidence'] <= 0.75:
					addto_doc(document, sentance.getvalue())
					addto_doc(document, word['word'], color=True)
					sentance = StringIO()
					sentance.write(" ")
				else:
					sentance.write("{0} ".format(word['word']))
				# print word['word']
			addto_doc(document, sentance.getvalue(), _break=True)
	doc.save('result.docx')

"""Adds to the running MS word doc. Flags are used to indicate formatting options.  
"""
def addto_doc(doc, sentance, _break=False, color=False, style=False):
	if (color):
		s = doc.add_run(sentance)
		font = s.font
		font.color.rgb = RGBColor(0xff, 0x00, 0x00)
		return
	if (style):
		s = doc.add_run(sentance)
		s.bold = True
		font = s.font
		font.color.rgb = RGBColor(0x46, 0x82, 0xb4)
		return
	r = doc.add_run(sentance)
	if(_break):
		r.add_break()
		r.add_break()

if __name__ == '__main__':
	api_key = input("Enter your API Key: ")
	transaction_id = input("Enter the transaction ID: ")
	main(api_key, transaction_id)
